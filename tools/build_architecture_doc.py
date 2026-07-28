from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
OUTPUT_PATH = OUTPUT_DIR / "RAG_Architecture_Design.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "667085"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "C8D0DA"
RISK = "9B1C1C"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Column widths must sum to {CONTENT_WIDTH_DXA}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_numbering_definition(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(lvl_text)

    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, num_id)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        set_run_font(paragraph.add_run(text))
    return paragraph


def add_callout(doc, label, text, color=DARK_BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.10
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT)
    p_pr.append(shd)
    lead = paragraph.add_run(f"{label}: ")
    set_run_font(lead, color=color, bold=True)
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_diagram(doc, text, caption):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F8FAFC")
    p_pr.append(shd)
    run = paragraph.add_run(text)
    set_run_font(run, name="Consolas", size=8.5, color=INK)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(8)
    set_run_font(cap.add_run(caption), size=9, color=MUTED, italic=True)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    mark_header_row(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, LIGHT_GRAY)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(paragraph.add_run(value), size=9.5, color=INK, bold=True)

    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            set_run_font(paragraph.add_run(str(value)), size=9.25)

    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)
    return table


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.font.size = Pt(12.5)
    subtitle.font.color.rgb = RGBColor.from_string(MUTED)
    subtitle.paragraph_format.space_after = Pt(16)

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    set_run_font(
        p.add_run("RAG BACKEND  |  ARCHITECTURE DESIGN"),
        size=8.5,
        color=MUTED,
        bold=True,
    )

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    set_run_font(fp.add_run("Trang "), size=9, color=MUTED)
    add_page_field(fp)


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    bullet_num = add_numbering_definition(doc, "bullet")
    decimal_num = add_numbering_definition(doc, "decimal")

    # First-page memo masthead.
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    set_run_font(kicker.add_run("TECHNICAL DESIGN"), size=10, color=BLUE, bold=True)
    doc.add_paragraph("Thiết kế kiến trúc hệ thống RAG", style="Title")
    doc.add_paragraph(
        "FastAPI, PostgreSQL và pgvector - phiên bản sau đợt khắc phục lỗi kỹ thuật",
        style="Subtitle",
    )
    for label, value in (
        ("Dự án", "FastAPI RAG Backend with PostgreSQL pgvector"),
        ("Phiên bản tài liệu", "1.1"),
        ("Ngày cập nhật", date(2026, 7, 28).strftime("%d/%m/%Y")),
        ("Phạm vi", "Kiến trúc hiện hành, quyết định thiết kế và lộ trình production"),
        ("Trạng thái", "Đã đồng bộ với mã nguồn và bộ kiểm thử 21 trường hợp"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(f"{label}: "), bold=True)
        set_run_font(p.add_run(value))

    add_callout(
        doc,
        "Kết luận",
        "Hệ thống là một modular monolith phù hợp cho MVP. Các lỗi về cấu hình database, "
        "validation, tính nguyên tử khi ingestion, lọc hybrid search, index truy vấn, provider "
        "và Docker đã được xử lý. Authentication và hàng đợi nền vẫn là backlog "
        "trước khi triển khai production.",
    )

    doc.add_heading("1. Mục tiêu và phạm vi", level=1)
    add_body(
        doc,
        "Tài liệu này mô tả kiến trúc logic, luồng dữ liệu, mô hình lưu trữ, giao diện API, "
        "cấu hình triển khai và các yêu cầu vận hành của backend RAG hiện tại."
    )
    add_list_item(doc, "Nạp tài liệu PDF, DOCX, TXT, Markdown, JSON và CSV.", bullet_num)
    add_list_item(doc, "Chia văn bản thành các đoạn chồng lấn và tạo embedding 1.536 chiều.", bullet_num)
    add_list_item(doc, "Truy hồi vector hoặc hybrid rồi tạo prompt có trích dẫn nguồn.", bullet_num)
    add_list_item(doc, "Sinh phản hồi qua OpenAI, Gemini hoặc mock provider.", bullet_num)
    add_list_item(doc, "Cung cấp REST API và triển khai cục bộ bằng Docker Compose.", bullet_num)

    doc.add_heading("Ngoài phạm vi hiện tại", level=2)
    add_body(
        doc,
        "Frontend, quản lý người dùng, phân quyền theo tenant, OCR tài liệu scan, reranker "
        "chuyên dụng và đánh giá chất lượng câu trả lời chưa nằm trong implementation hiện hành."
    )

    doc.add_heading("2. Bối cảnh hệ thống", level=1)
    add_diagram(
        doc,
        """
+------------------+       HTTPS/JSON       +-----------------------------+
| Client / Frontend| ---------------------> | FastAPI REST API             |
+------------------+                        | /api/v1                      |
                                            +--------------+--------------+
                                                           |
                         +---------------------------------+---------------------------------+
                         |                                                                   |
                         v                                                                   v
              +----------------------+                                           +----------------------+
              | PostgreSQL + pgvector|                                           | AI Providers         |
              | documents / chunks   |                                           | OpenAI/Gemini/Mock   |
              +----------------------+                                           +----------------------+
""".strip(),
        "Hình 1 - Sơ đồ bối cảnh hệ thống",
    )
    add_body(
        doc,
        "FastAPI là biên giao tiếp duy nhất. PostgreSQL lưu metadata, nội dung chunk, full-text "
        "document và embedding; dịch vụ AI bên ngoài chỉ được gọi qua lớp provider."
    )

    doc.add_heading("3. Kiến trúc logic", level=1)
    add_diagram(
        doc,
        """
[main.py / lifespan / CORS]
             |
             v
[API schemas] ---> [API routes + orchestration]
                          |
          +---------------+----------------+----------------+
          |               |                |                |
          v               v                v                v
      [Parser]         [Chunker]       [Embedding]     [Prompt + LLM]
          |               |                |                ^
          +---------------+-------> [Vector Store] ---------+
                                         |
                                         v
                              [SQLAlchemy + PostgreSQL]
""".strip(),
        "Hình 2 - Kiến trúc component của modular monolith",
    )

    add_table(
        doc,
        ["Thành phần", "Trách nhiệm", "Tệp chính"],
        [
            ("Bootstrap", "Khởi tạo ứng dụng, lifespan và CORS.", "main.py"),
            ("API", "Validation request/response và điều phối use case.", "app/api/"),
            ("Parser", "Trích xuất văn bản từ định dạng được hỗ trợ.", "services/parser.py"),
            ("Chunker", "Chia đoạn theo ký tự, kiểm soát overlap.", "services/chunker.py"),
            ("Embedding", "Adapter OpenAI, Gemini và mock.", "services/embedding.py"),
            ("Retrieval", "Vector search, FTS, RRF và lưu chunk.", "services/vector_store.py"),
            ("Generation", "Xây prompt có citation và gọi LLM.", "prompt_builder.py; llm.py"),
            ("Persistence", "Session, schema, pgvector và index.", "database.py; models.py"),
        ],
        [1800, 4680, 2880],
    )

    add_callout(
        doc,
        "Ranh giới kiến trúc",
        "Routes hiện là application orchestrator. Khi use case tăng, nên tách thành "
        "IngestionService và QueryService để API chỉ còn nhiệm vụ chuyển đổi giao thức.",
    )

    doc.add_heading("4. Luồng ingestion", level=1)
    add_diagram(
        doc,
        """
Upload/Text -> Validate type, size, chunk window -> Parse -> Split chunks
                                                     |
                                                     v
                                          Embed every chunk
                                                     |
                                                     v
                             One DB transaction: Document + DocumentChunk[]
                                                     |
                                                     v
                                             Commit / Rollback
""".strip(),
        "Hình 3 - Luồng nạp tài liệu",
    )
    add_list_item(doc, "Kiểm tra extension, giới hạn mặc định 20 MB và tham số chunk.", decimal_num)
    add_list_item(doc, "Parse nội dung trước khi mở transaction ghi database.", decimal_num)
    add_list_item(doc, "Tạo embedding; lỗi provider không được âm thầm chuyển sang mock.", decimal_num)
    add_list_item(doc, "Ghi Document và toàn bộ chunks trong một transaction.", decimal_num)
    add_list_item(doc, "Rollback khi bất kỳ bước ghi dữ liệu nào thất bại.", decimal_num)
    add_callout(
        doc,
        "Giới hạn",
        "Embedding vẫn chạy tuần tự trong request worker. Tài liệu lớn cần batch embedding và "
        "background queue để tránh timeout; đây là bước mở rộng production, không phải lỗi "
        "tính đúng của phiên bản hiện tại.",
    )

    doc.add_heading("5. Luồng retrieval và generation", level=1)
    add_diagram(
        doc,
        """
Question -> Query embedding
              |
              +--> Vector search (cosine, HNSW) --------+
              |                                         |
              +--> Full-text search (simple, GIN) ------+--> RRF (k=60)
                                                              |
                                                              v
                                                       Top-K contexts
                                                              |
                                                              v
                                             Prompt with [Source #n] -> LLM
""".strip(),
        "Hình 4 - Luồng truy vấn RAG hybrid",
    )
    add_body(
        doc,
        "Vector search sử dụng cosine distance của pgvector. Hybrid search kết hợp thứ hạng "
        "vector và PostgreSQL full-text search bằng Reciprocal Rank Fusion. Cấu hình text search "
        "'simple' giữ token phù hợp hơn cho nội dung tiếng Việt so với cấu hình 'english'."
    )
    add_callout(
        doc,
        "Lưu ý thuật ngữ",
        "Cơ chế full-text hiện tại dùng ts_rank_cd; không được mô tả là BM25. Nếu yêu cầu BM25 "
        "thực sự, cần bổ sung extension hoặc search engine chuyên dụng và đo lại chất lượng.",
    )

    doc.add_heading("6. Thiết kế dữ liệu", level=1)
    add_diagram(
        doc,
        """
Document (1) --------------------------------------< (N) DocumentChunk
 id: UUID string                                         id: UUID string
 filename, file_type, file_size                          document_id: FK
 chunk_count, created_at                                 chunk_index, content
                                                         metadata_json
                                                         embedding vector(1536)
""".strip(),
        "Hình 5 - Mô hình quan hệ dữ liệu",
    )
    add_body(
        doc,
        "DocumentChunk phụ thuộc Document và bị xóa cascade. Embedding không nullable ở schema "
        "mới; metadata dùng default callable để tránh chia sẻ mutable object."
    )
    add_table(
        doc,
        ["Index", "Loại", "Mục đích"],
        [
            ("ix_document_chunks_document_id", "B-tree", "Lọc chunk theo tài liệu."),
            ("ix_document_chunks_embedding_hnsw", "HNSW cosine", "Approximate nearest-neighbor search."),
            ("ix_document_chunks_fts_simple", "GIN expression", "Full-text retrieval cấu hình simple."),
        ],
        [3400, 2000, 3960],
    )
    add_callout(
        doc,
        "Tính tương thích embedding",
        "Không trộn vector từ nhiều model trong cùng collection. Khi đổi provider hoặc model, "
        "cần re-index toàn bộ tài liệu hoặc bổ sung trường embedding_model/collection.",
    )

    doc.add_heading("7. Thiết kế API", level=1)
    add_table(
        doc,
        ["Method", "Endpoint", "Mục đích"],
        [
            ("GET", "/api/v1/health", "Kiểm tra DB và pgvector."),
            ("POST", "/api/v1/documents/upload", "Upload và index tài liệu."),
            ("POST", "/api/v1/documents/ingest-text", "Index raw text."),
            ("GET", "/api/v1/documents", "Liệt kê tài liệu."),
            ("DELETE", "/api/v1/documents/{doc_id}", "Xóa tài liệu và chunks."),
            ("POST", "/api/v1/retrieval/search", "Vector hoặc hybrid search."),
            ("POST", "/api/v1/chat/query", "RAG query hoàn chỉnh."),
        ],
        [1250, 3970, 4140],
    )
    add_body(
        doc,
        "Các request sử dụng Pydantic để ràng buộc search_type, top_k, chunk_size và overlap. "
        "Cả vector và hybrid search đều hỗ trợ document_id filter. OpenAPI có tại /docs."
    )
    add_callout(
        doc,
        "Bảo mật API",
        "CORS mặc định chỉ cho phép localhost:3000 và localhost:5173. Hệ thống chưa có "
        "authentication, authorization hoặc rate limiting; không mở API ra Internet trước khi "
        "bổ sung các lớp kiểm soát này.",
        color=RISK,
    )

    doc.add_heading("8. Cấu hình và triển khai", level=1)
    add_body(
        doc,
        "Chạy toàn bộ stack bằng docker compose up --build. Compose khởi tạo PostgreSQL pgvector "
        "và chỉ khởi động API sau khi database health check thành công."
    )
    add_table(
        doc,
        ["Ngữ cảnh", "DATABASE_URL", "Giải thích"],
        [
            (
                "API chạy trên máy host",
                "postgresql+psycopg://...@localhost:45432/rag_db",
                "Compose publish cổng 45432 ra host.",
            ),
            (
                "API chạy trong Compose",
                "postgresql+psycopg://...@postgres:5432/rag_db",
                "Dùng service name và cổng nội bộ.",
            ),
        ],
        [2100, 4380, 2880],
    )
    add_body(
        doc,
        "Biến môi trường mẫu nằm trong .env.example. Provider mặc định là mock; khi chọn OpenAI "
        "hoặc Gemini phải cung cấp API key tương ứng. Model name được cấu hình độc lập."
    )

    doc.add_heading("9. Quyết định thiết kế", level=1)
    add_table(
        doc,
        ["Quyết định", "Lý do", "Hệ quả"],
        [
            ("Modular monolith", "Đơn giản cho MVP và dễ debug.", "Cần tách worker khi tải ingestion tăng."),
            ("PostgreSQL làm unified store", "Giảm số hệ thống vận hành.", "FTS không tương đương BM25 chuyên dụng."),
            ("HNSW cosine", "Latency tốt cho truy hồi gần đúng.", "Tốn bộ nhớ và thời gian build index."),
            ("RRF cho hybrid", "Không cần chuẩn hóa hai thang điểm.", "Cần tuning top_k và đánh giá offline."),
            ("Provider adapter", "Đổi nhà cung cấp qua cấu hình.", "Phải giữ dimension/model đồng nhất."),
            ("Fail fast khi thiếu key", "Tránh vô tình chạy mock.", "Lỗi cấu hình xuất hiện sớm và rõ."),
        ],
        [2500, 3440, 3420],
    )

    doc.add_heading("10. Chất lượng, vận hành và backlog", level=1)
    doc.add_heading("Đã hoàn thành", level=2)
    add_list_item(doc, "21 unit tests đạt; kiểm tra compile và OpenAPI đạt.", bullet_num)
    add_list_item(doc, "Docker Compose được validate và có đầy đủ API + database service.", bullet_num)
    add_list_item(doc, "Ingestion atomic, validation biên và health check pgvector.", bullet_num)
    add_list_item(doc, "HNSW, GIN và document_id indexes được khởi tạo idempotent.", bullet_num)
    add_list_item(doc, "Thông báo mock đã sửa UTF-8 và không bịa nguồn khi context rỗng.", bullet_num)

    doc.add_heading("Backlog trước production", level=2)
    add_list_item(doc, "P0 - Authentication, authorization theo document/tenant và rate limiting.", bullet_num)
    add_list_item(doc, "P0 - Quy trình backup và diễn tập rollback database migration.", bullet_num)
    add_list_item(doc, "P1 - Background worker, batch embedding, retry có backoff và idempotency key.", bullet_num)
    add_list_item(doc, "P1 - Structured logging, metrics, tracing và cảnh báo lỗi provider.", bullet_num)
    add_list_item(doc, "P1 - Integration test với PostgreSQL/pgvector và contract test provider.", bullet_num)
    add_list_item(doc, "P2 - Reranker, evaluation dataset, recall@k, groundedness và latency SLO.", bullet_num)
    add_list_item(doc, "P2 - OCR, antivirus scan và lưu object storage cho file gốc.", bullet_num)

    doc.add_heading("11. Tiêu chí nghiệm thu kiến trúc", level=1)
    add_list_item(doc, "docker compose up --build đưa database và API về trạng thái healthy.", decimal_num)
    add_list_item(doc, "Upload sai loại, quá kích thước hoặc chunk window sai bị từ chối.", decimal_num)
    add_list_item(doc, "Lỗi embedding không để lại document/chunk ghi dở.", decimal_num)
    add_list_item(doc, "Vector và hybrid search tôn trọng document_id filter.", decimal_num)
    add_list_item(doc, "Câu trả lời RAG trả kèm retrieved_contexts và citation trong prompt.", decimal_num)
    add_list_item(doc, "Đổi provider thiếu API key thất bại rõ ràng, không fallback mock.", decimal_num)
    add_list_item(doc, "Bộ test và OpenAPI validation tiếp tục đạt trong CI.", decimal_num)

    doc.add_heading("12. Bản đồ mã nguồn", level=1)
    add_body(doc, "Các tệp kiến trúc quan trọng:", bold_lead="Các tệp kiến trúc quan trọng:")
    for path, role in (
        ("main.py", "bootstrap, lifespan và CORS"),
        ("app/config.py", "cấu hình từ environment"),
        ("app/api/routes.py", "REST endpoints và orchestration"),
        ("app/api/schemas.py", "API contracts và validation"),
        ("app/models.py", "mô hình dữ liệu và HNSW index"),
        ("app/database.py", "session, pgvector và idempotent indexes"),
        ("app/services/", "parser, chunker, provider, retrieval và generation"),
        ("docker-compose.yml", "topology triển khai local"),
        ("tests/test_rag_pipeline.py", "kiểm thử pipeline và lỗi biên"),
    ):
        add_list_item(doc, f"{path} - {role}.", bullet_num)

    # Prevent an accidental trailing blank page.
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True

    doc.core_properties.title = "Thiết kế kiến trúc hệ thống RAG"
    doc.core_properties.subject = "FastAPI, PostgreSQL và pgvector"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "RAG, FastAPI, PostgreSQL, pgvector, architecture"
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
