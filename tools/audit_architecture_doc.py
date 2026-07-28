from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "RAG_Architecture_Design.docx"
EXPECTED_WIDTH = 9360
EXPECTED_INDENT = 120


def int_attr(element, name):
    return int(element.get(qn(name)))


def main():
    doc = Document(DOCX)
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.50
    assert round(section.page_height.inches, 2) == 11.00
    assert all(
        round(value.inches, 2) == 1.00
        for value in (
            section.top_margin,
            section.right_margin,
            section.bottom_margin,
            section.left_margin,
        )
    )
    assert round(section.header_distance.inches, 3) == 0.492
    assert round(section.footer_distance.inches, 3) == 0.492

    style_expectations = {
        "Normal": (11, 0, 6),
        "Heading 1": (16, 16, 8),
        "Heading 2": (13, 12, 6),
        "Heading 3": (12, 8, 4),
    }
    for style_name, (size, before, after) in style_expectations.items():
        style = doc.styles[style_name]
        assert round(style.font.size.pt, 1) == size
        assert round(style.paragraph_format.space_before.pt, 1) == before
        assert round(style.paragraph_format.space_after.pt, 1) == after

    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        assert int_attr(tbl_w, "w:w") == EXPECTED_WIDTH
        assert int_attr(tbl_ind, "w:w") == EXPECTED_INDENT

        grid_widths = [
            int_attr(node, "w:w")
            for node in table._tbl.tblGrid.findall(qn("w:gridCol"))
        ]
        assert sum(grid_widths) == EXPECTED_WIDTH
        for row in table.rows:
            cell_widths = [
                int_attr(cell._tc.get_or_add_tcPr().find(qn("w:tcW")), "w:w")
                for cell in row.cells
            ]
            assert cell_widths == grid_widths

    numbered_paragraphs = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph._p.get_or_add_pPr().find(qn("w:numPr")) is not None
    ]
    assert len(numbered_paragraphs) >= 20

    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "PLACEHOLDER" not in all_text
    assert ":codex-" not in all_text
    assert "BM25 thực sự" in all_text
    assert "16 unit tests" in all_text
    assert len(doc.tables) == 5

    print(
        f"Structural audit: OK | paragraphs={len(doc.paragraphs)} "
        f"| tables={len(doc.tables)} | numbered={len(numbered_paragraphs)}"
    )


if __name__ == "__main__":
    main()
